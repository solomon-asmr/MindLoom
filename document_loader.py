import os
import csv
import fitz
from docx import Document

def _extract_images_from_pdf(file_path):
    """Extract all images from a PDF file.
    
    Returns:
        list of image file paths
    """
    doc = fitz.open(file_path)
    image_paths = []
    
    for page_num in range(len(doc)):
        page = doc[page_num]
        images = page.get_images(full=True)
        
        for img_index, img in enumerate(images):
            xref = img[0]
            
            try:
                base_image = doc.extract_image(xref)
                image_bytes = base_image["image"]
                image_ext = base_image["ext"]
                
                # Save temporarily
                image_path = file_path.replace(
                    os.path.splitext(file_path)[1],
                    f"_page{page_num}_img{img_index}.{image_ext}"
                )
                
                with open(image_path, "wb") as f:
                    f.write(image_bytes)
                
                image_paths.append(image_path)
            except Exception:
                continue
    
    doc.close()
    return image_paths


def _extract_images_from_docx(file_path):
    """Extract all images from a DOCX file.
    
    Returns:
        list of image file paths
    """
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
    
    doc = Document(file_path)
    image_paths = []
    
    for i, rel in enumerate(doc.part.rels.values()):
        if "image" in rel.reltype:
            try:
                image_data = rel.target_part.blob
                
                # Detect extension from content type
                content_type = rel.target_part.content_type
                if "png" in content_type:
                    ext = "png"
                elif "gif" in content_type:
                    ext = "gif"
                elif "webp" in content_type:
                    ext = "webp"
                else:
                    ext = "jpg"
                
                image_path = file_path.replace(
                    os.path.splitext(file_path)[1],
                    f"_img{i}.{ext}"
                )
                
                with open(image_path, "wb") as f:
                    f.write(image_data)
                
                image_paths.append(image_path)
            except Exception:
                continue
    
    return image_paths

def load_pdf(file_path, analyze_images_fn=None):
    """Extract text from a PDF file."""
    doc = fitz.open(file_path)
    
    text = ""
    for page in doc:
        text += page.get_text() + "\n"
    
    doc.close()
    # Extract and analyze images if function provided
    if analyze_images_fn:
        image_paths = _extract_images_from_pdf(file_path)
        for i, img_path in enumerate(image_paths):
            try:
                result = analyze_images_fn(img_path)
                if result["success"]:
                    text += f"\n\n[Image {i+1} in document]: {result['text']}\n"
            except Exception:
                pass
            finally:
                if os.path.exists(img_path):
                    os.remove(img_path)
    
    return text.strip()

def load_txt(file_path):
    """Read a plain text file."""
    with open(file_path, "r", encoding="utf-8") as f:
        return f.read().strip()


def load_docx(file_path, analyze_images_fn=None):
    """Extract text from a Word document."""
    doc = Document(file_path)
    
    text = ""
    for paragraph in doc.paragraphs:
        text += paragraph.text + "\n"
    
    if analyze_images_fn:
        image_paths = _extract_images_from_docx(file_path)
        for i, img_path in enumerate(image_paths):
            try:
                result = analyze_images_fn(img_path)
                if result["success"]:
                    text += f"\n\n[Image {i+1} in document]: {result['text']}\n"
            except Exception:
                pass
            finally:
                if os.path.exists(img_path):
                    os.remove(img_path)
    
    return text.strip()


def load_csv(file_path):
    """Extract text from a CSV file."""
    text = ""
    
    with open(file_path, "r", encoding="utf-8") as f:
        reader = csv.reader(f)
        for row in reader:
            text += " ".join(row) + "\n"
    
    return text.strip()


LOADERS = {
    ".pdf": load_pdf,
    ".txt": load_txt,
    ".docx": load_docx,
    ".csv": load_csv,
}


def load_document(file_path, analyze_images_fn=None):
    """Auto-detect file type and extract text."""
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext not in LOADERS:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {list(LOADERS.keys())}")
    
    # Pass image analyzer to PDF and DOCX loaders
    if ext in [".pdf", ".docx"] and analyze_images_fn:
        text = LOADERS[ext](file_path, analyze_images_fn=analyze_images_fn)
    else:
        text = LOADERS[ext](file_path)
    
    if not text:
        raise ValueError(f"No text could be extracted from {file_path}")
    
    return text


def get_supported_extensions():
    """Return list of supported file extensions."""
    return list(LOADERS.keys())